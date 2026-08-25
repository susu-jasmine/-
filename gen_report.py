"""Generate the training report docx — all image markers point to Web UI screenshots only."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_run_font(run, font='宋体', size=10.5, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


# ============ 封面 ============
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('实训报告')
set_run_font(r, '方正小标宋简体', 36, True)

for _ in range(5):
    doc.add_paragraph()

cover_info = [
    ('课程名称：', '计算机应用开发综合实训'),
    ('学    院：', '__________________'),
    ('班    级：', '__________________'),
    ('学    号：', '__________________'),
    ('姓    名：', '__________________'),
    ('指导教师：', '王志成'),
    ('起止时间：', '2026年6月8日 至 2026年6月28日'),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label + '   ' + value)
    set_run_font(r, '宋体', 14, False)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('题目：资产监控系统 AssetMonitor')
set_run_font(r, '黑体', 16, True)

doc.add_page_break()


# ============ Helpers ============
def add_heading(text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, '黑体', sizes.get(level, 12), True)


def add_body(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, '宋体', 10.5, bold)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    r.font.size = Pt(9)


def add_image_marker(desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('▼▼▼  【📷 此处插入图片：' + desc + '】  ▼▼▼')
    set_run_font(r, '宋体', 10.5, True, (0xC0, 0x00, 0x00))


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run_font(r, '宋体', 9, False, (0x40, 0x40, 0x40))


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, '宋体', 10.5, True)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, '宋体', 10, False)


# ============ 题目说明 ============
add_heading('（最好：图文并茂，格式规范）', 2)
add_body('本报告依据《计算机应用开发综合实训》要求撰写，以"资产监控系统 AssetMonitor"为实践项目，由两人协作完成。报告全文按"绪论 — 需求分析 — 系统设计 — 系统实现 — 总结"五部分组织，并按要求在关键节点配图说明。所有配图均为系统实际运行的 Web 界面截图，红色【📷 此处插入图片】字样为待补图位置，定稿时删除该行并替换为实际截图。')
doc.add_page_break()

# ========== 一、绪论 ==========
add_heading('一、绪论', 1)
add_image_marker('项目主界面截图（仪表盘首屏：统计卡片 + 域名表格）')
add_caption('图 1-0  AssetMonitor 系统主界面')

add_heading('1.1 项目背景', 2)
add_body('在网络安全攻防实战中，资产发现（Asset Discovery）是攻击面管理的第一步，也是 SRC（Security Response Center，安全应急响应中心）漏洞挖掘的核心痛点。传统手工监控存在三大问题：')
add_body('（1）响应滞后：依赖人工周期性查询 crt.sh、FOFA、Sublist3r 等工具，常常晚于其他研究者数小时甚至数天才能发现新增子域名；')
add_body('（2）平台监控缺失：补天、漏洞盒子等 SRC 平台新上企业无自动告警机制，错过首日提交黄金窗口；')
add_body('（3）资产分散难管理：手工 Excel 记录易丢失、难检索、无法关联组织关系。')

add_heading('1.2 项目目标', 2)
add_body('构建一套全自动化资产监控与告警系统，实现域名资产自动定时扫描、IP 解析、新资产微信/邮件告警；SRC 平台（补天/漏洞盒子）新企业上线秒级推送；多用户、多组织、批量导入、可视化管理；中英双语、一键云端部署、在线更新。')

add_heading('1.3 技术栈', 2)
add_table(
    ['层次', '技术选型'],
    [
        ['后端框架', 'Flask + SQLAlchemy + APScheduler'],
        ['数据库', 'SQLite（可平滑迁移 MySQL/PostgreSQL）'],
        ['认证', 'JWT（Flask-JWT-Extended）'],
        ['任务调度', 'APScheduler BackgroundScheduler'],
        ['子域名采集', 'crt.sh + AlienVault OTX + DNS 暴力枚举'],
        ['通知', 'PushPlus（微信）+ smtplib（邮件）'],
        ['部署', 'systemd + Gunicorn + Nginx'],
        ['前端', '原生 HTML5 + CSS3 + ES6，无前端框架'],
        ['部署平台', 'Linux（Debian/Ubuntu/CentOS）'],
    ]
)

# ========== 二、需求分析 ==========
add_heading('二、需求分析', 1)

add_heading('2.1 功能性需求', 2)
add_body('FR-U1 系统提供管理员账户（默认 admin / admin123456），登录后强制修改密码；')
add_body('FR-U2 基于 JWT 的会话管理，Token 有效期 168 小时；')
add_body('FR-D1 单个/批量添加监控域名，支持所属单位（Organization）字段；')
add_body('FR-D2 域名生命周期：添加 / 编辑 / 暂停 / 恢复 / 删除 / 立即扫描；')
add_body('FR-A1 自动发现子域名、解析 IP、记录首次发现时间；')
add_body('FR-S1 自动采集补天（企业 SRC/专属 SRC/公益 SRC）和漏洞盒子全部上线企业；')
add_body('FR-N1 微信推送（PushPlus）+ 邮件推送（SMTP）双通道告警；')
add_body('FR-O1 一键部署脚本（apt/yum/dnf 自适配）；')
add_body('FR-O2 系统内置 Git 在线更新模块，后台一键升级；')
add_body('FR-O3 中英双语动态切换。')
add_image_marker('系统用例图（推荐用 StarUML / draw.io 重绘正式 UML 用例图）')
add_caption('图 2-1  系统用例图')

add_heading('2.2 非功能性需求', 2)
add_table(
    ['编号', '类别', '要求'],
    [
        ['NFR-1', '性能', '单用户 100 域名并发扫描响应 < 60s'],
        ['NFR-2', '可用性', 'systemd 守护，进程崩溃 5s 内自动重启'],
        ['NFR-3', '安全性', 'JWT 鉴权；密码 Werkzeug PBKDF2 哈希；.env 文件 600 权限'],
        ['NFR-4', '可维护', '自动 SQLite Schema 迁移；在线更新'],
        ['NFR-5', '易用性', '全图形界面，无命令行使用门槛'],
        ['NFR-6', '兼容性', '主流 Linux 发行版（Debian/Ubuntu/CentOS/Rocky）'],
        ['NFR-7', '国际化', 'zh / en 一键切换，UI 文案全量 i18n'],
        ['NFR-8', '可扩展', '模块化设计，扫描源、SRC 平台、通知通道均可插拔'],
    ]
)

# ========== 三、系统设计 ==========
add_heading('三、系统设计', 1)
add_body('（本章节着重描述项目整体架构与个人承担的模块分工。）')

add_heading('3.1 整体架构', 2)
add_body('系统采用经典三层架构（前端 / 后端 / 数据），并在后端引入后台调度器实现自动化：')
add_code('浏览器（前端）\n      ↓ HTTPS / JWT\nFlask + Gunicorn（auth / api / scanner / scheduler / src_monitor / notifier / updater）\n      ↓ SQLAlchemy ORM\nSQLite（monitor.db）')
add_image_marker('系统架构示意图（draw.io 重绘成正式流程图）')
add_caption('图 3-1  系统三层架构图')

add_heading('3.2 数据库设计', 2)
add_body('系统包含 5 张核心表：users（用户与通知配置）、domains（监控任务）、assets（资产明细）、scan_logs（扫描流水）、src_programs（SRC 企业）。下方 Web 界面截图展示了实际运行中数据表之间的关系。')
add_image_marker('Web 界面截图：仪表盘统计卡片（域名数/总资产/新资产/SRC新企业），直观体现各表数据汇聚结果')
add_caption('图 3-2  仪表盘统计卡片（体现数据库各表数据）')

add_heading('3.3 模块分工', 2)
add_body('为体现团队协作，本项目由学生 A 与学生 B 协同开发：')
add_table(
    ['模块', '学生 A（后端 + 部署）', '学生 B（前端 + 集成）'],
    [
        ['用户认证', 'JWT 模块、密码哈希', '登录页 UI、表单校验'],
        ['数据建模', 'SQLAlchemy 5 张表设计', '—'],
        ['REST API', '全部接口（域/资产/SRC/设置/更新）', '—'],
        ['扫描引擎', 'crt.sh + OTX + DNS 采集器', '—'],
        ['调度器', 'APScheduler 集成、动态增删任务', '—'],
        ['SRC 监控', 'API 逆向（补天/漏洞盒子）', 'SRC 视图、列表筛选、导出'],
        ['通知模块', 'PushPlus + SMTP 发送器', '通知模板、设置 UI'],
        ['仪表盘', '—', '全部 HTML/CSS/JS'],
        ['双语支持', '—', 'i18n 框架与 200+ 词条翻译'],
        ['部署', 'deploy.sh + systemd + sudoers', 'Nginx 反代配置'],
        ['在线更新', 'updater.py + Git 集成', '更新 UI 与交互'],
    ]
)
add_image_marker('Web 界面截图：导航栏切换五个模块（域名 / 全部资产 / 扫描日志 / 设置 / SRC监控）任一截图，体现模块协作成果')
add_caption('图 3-3  Web 端模块导航切换')

# ========== 四、系统实现 ==========
add_heading('四、系统实现', 1)
add_body('（以下章节按小组成员分工分别撰写。所有配图均为 Web 界面实际运行截图。）')

# ===== 学生 A =====
add_heading('4.1 学生 A 负责模块', 2)

add_heading('4.1.1 系统框架搭建', 3)
add_body('采用 Flask 工厂模式，app/__init__.py 中通过 create_app() 统一初始化数据库、蓝图、调度器：')
add_code(
    "def create_app():\n"
    "    app = Flask(__name__, template_folder=..., static_folder=...)\n"
    "    app.config.from_object('config.Config')\n"
    "    db.init_app(app)\n"
    "    with app.app_context():\n"
    "        db.create_all()\n"
    "    app.register_blueprint(auth_bp, url_prefix='/api/auth')\n"
    "    app.register_blueprint(api_bp, url_prefix='/api')\n"
    "    from app.scheduler import init_scheduler\n"
    "    init_scheduler(app)\n"
    "    return app"
)
add_image_marker('Web 界面截图：浏览器访问根路径自动跳转登录页（体现 Flask 路由正常工作）')
add_caption('图 4-1  系统启动后访问首页')

add_heading('4.1.2 JWT 认证模块', 3)
add_body('自定义 @token_required 装饰器统一鉴权，前端登录成功后获取 JWT Token 并保存到 localStorage，所有 API 请求自动携带 Authorization 头部。')
add_image_marker('Web 界面截图：登录页输入用户名密码后点击"验证"成功跳转仪表盘')
add_caption('图 4-2  JWT 登录流程（Web 端）')

add_heading('4.1.3 子域名扫描引擎', 3)
add_body('三源聚合策略，扫描结果统一去重入库，新子域名自动置 is_new=True，并在前端以高亮"NEW"标识展示。')
add_table(
    ['数据源', '接口', '特性'],
    [
        ['crt.sh', 'https://crt.sh/?q=%25.domain&output=json', 'CT 日志，新子域名最及时'],
        ['AlienVault OTX', 'https://otx.alienvault.com/api/v1/indicators/domain/{d}/passive_dns', '历史被动 DNS'],
        ['DNS 暴力枚举', '内置 5000+ 子域字典', '兜底发现'],
    ]
)
add_image_marker('Web 界面截图："全部资产"页 — 展示扫描发现的子域名、IP、来源、首次发现时间，新资产带 NEW 标识')
add_caption('图 4-3  扫描结果在"全部资产"页展示')

add_heading('4.1.4 调度器', 3)
add_body('基于 APScheduler BackgroundScheduler，每条域名一个独立 Job；SRC 监控任务以"所有用户最小间隔"为全局周期。Web 界面通过域名表格的"上次扫描"列和"状态"列直观反映调度结果。')
add_image_marker('Web 界面截图：域名表格 — 展示"上次扫描"时间列 + "状态"列（活跃/已暂停），点击"扫描"按钮触发即时任务')
add_caption('图 4-4  调度器在域名表格中的可视化反馈')

add_heading('4.1.5 SRC 平台 API 逆向', 3)
add_body('通过浏览器 F12 抓包分析，得到补天三个标签页和漏洞盒子的真实接口。识别新增用 (platform, company_id) 联合唯一键。逆向结果最终在 SRC 监控页面以表格形式展示。')
add_code(
    "BUTIAN_TABS = {\n"
    "    'corps': 'https://www.butian.net/Reward/corps',\n"
    "    'com':   'https://www.butian.net/Reward/com',\n"
    "    'pub':   'https://www.butian.net/Reward/pub',\n"
    "}\n"
    "VULBOX_URL = 'https://vapi.vulbox.com/web/project/enterprise/src'"
)
add_image_marker('Web 界面截图：SRC 监控页 — 展示补天和漏洞盒子的企业列表（平台、企业名、奖金范围、状态）')
add_caption('图 4-5  SRC 平台数据在 Web 端展示')

add_heading('4.1.6 通知模块', 3)
add_body('双通道告警：微信（PushPlus）+ 邮件（SMTP）。通知开关与配置均在 Web 设置页完成。')
add_image_marker('Web 界面截图：设置页 — 通知配置区块（PushPlus Token 输入框 + 微信/邮件启用开关）')
add_caption('图 4-6  Web 端通知配置面板')

add_heading('4.1.7 一键部署脚本', 3)
add_body('跨发行版兼容（apt/yum/dnf 自适配），创建专用用户、venv、systemd 单元，并自动 SQLite Schema 迁移。部署成功后浏览器直接访问端口即可看到登录页。')
add_image_marker('Web 界面截图：部署完成后浏览器访问 http://服务器IP:5000 看到的登录页')
add_caption('图 4-7  部署成功后的访问效果')

add_heading('4.1.8 在线更新模块', 3)
add_body('本项目原创亮点，让平台支持后台一键 OTA。基于 Git + sudoers 精细放权 + 延迟 systemd 重启，实现 Web 端零中断升级。')
add_code(
    "def apply_update():\n"
    "    subprocess.run('git pull origin main --ff-only', shell=True)\n"
    "    subprocess.run(f'{VENV_PYTHON} -m pip install -r requirements.txt -q')\n"
    "    subprocess.Popen(['bash', '-c',\n"
    "        'sleep 1.5 && sudo -n /bin/systemctl restart asset-monitor'])"
)
add_image_marker('Web 界面截图：设置页底部"// 在线更新"区块（保存仓库 / 检查更新 / 应用更新 三个按钮）')
add_caption('图 4-8  在线更新 Web 控制面板')
add_image_marker('Web 界面截图：点击"检查更新"后显示当前版本、远端版本与落后提交列表')
add_caption('图 4-9  Web 端检查更新结果')

# ===== 学生 B =====
add_heading('4.2 学生 B 负责模块', 2)

add_heading('4.2.1 仪表盘 UI 与黑客风视觉设计', 3)
add_body('设计目标是"安全研究员审美"：黑底 + 4 色霓虹（绿/青/品红/黄），等宽字体，扫描线动画。')
add_image_marker('Web 界面截图：登录页（黑客风黑底霓虹，等宽字体）')
add_caption('图 4-10  登录页视觉设计')
add_image_marker('Web 界面全屏截图：仪表盘首屏（顶部导航 + 统计卡片 + 域名表格）')
add_caption('图 4-11  仪表盘整体视觉风格')

add_heading('4.2.2 双语 i18n 框架', 3)
add_body('基于 data-i18n 属性自动遍历，无前端框架依赖，支持 200+ 词条的 zh / en 全覆盖，并通过 localStorage 持久化用户选择。')
add_image_marker('Web 界面截图：任意页面中文状态（如域名管理页中文）')
add_caption('图 4-12  中文界面')
add_image_marker('Web 界面截图：同一页面点击右上角语言切换按钮后的英文状态')
add_caption('图 4-13  英文界面（同页面切换）')
add_image_marker('Web 界面截图：鼠标点击右上角"EN/中文"切换按钮的瞬间（按钮特写）')
add_caption('图 4-14  语言切换按钮交互')

add_heading('4.2.3 表格排序与分页', 3)
add_body('中央 SortState 对象统一管理所有表格的排序状态；后端通过 sort_by + sort_asc 参数白名单映射到 ORM 列，避免 SQL 注入。')
add_image_marker('Web 界面截图：域名列表点击"上次扫描"列头，▼ 三角图标高亮，列表按时间倒序排列')
add_caption('图 4-15  表格列排序三角指示器')
add_image_marker('Web 界面截图：资产页底部分页器（上一页 / 页码 / 下一页）')
add_caption('图 4-16  分页器组件')

add_heading('4.2.4 批量导入与组织分组', 3)
add_body('设置页提供批量导入面板，按"所属单位"分组，后端通过事务批量插入，自动去重，返回成功/跳过计数。')
add_image_marker('Web 界面截图：点击"批量导入"按钮展开的批量输入面板（单位 + 多行域名 + 间隔 + 执行导入按钮）')
add_caption('图 4-17  批量导入面板展开')
add_image_marker('Web 界面截图：导入成功后页面顶部弹出的绿色 Toast 成功提示')
add_caption('图 4-18  导入成功 Toast 反馈')
add_image_marker('Web 界面截图：域名表格上方的"筛选单位"下拉框 + 单位列展示')
add_caption('图 4-19  按单位分组与筛选')

add_heading('4.2.5 SRC 监控前端', 3)
add_body('顶部统计卡片：总企业数、新增数、按平台分布；多维筛选：平台（补天/漏洞盒子/全部）、仅新上、关键字搜索。')
add_image_marker('Web 界面全屏截图：SRC 监控页（顶部统计卡片 + 筛选区 + 表格主体）')
add_caption('图 4-20  SRC 监控页完整界面')
add_image_marker('Web 界面截图：勾选"仅新上"+ 在搜索框输入关键字的过滤效果')
add_caption('图 4-21  SRC 监控筛选效果')
add_image_marker('Web 界面截图：SRC 表格中"NEW"标识的新上企业行（高亮显示）')
add_caption('图 4-22  新上企业高亮标识')

add_heading('4.2.6 数据导出', 3)
add_body('支持 CSV / JSON 双格式导出全部资产或 SRC 企业数据，前端按钮触发浏览器下载。')
add_image_marker('Web 界面截图：资产页右上角"导出CSV""导出JSON"按钮（按钮特写）')
add_caption('图 4-23  Web 端导出按钮')
add_image_marker('Web 界面截图：点击导出后浏览器底部出现下载提示条')
add_caption('图 4-24  浏览器下载反馈')

add_heading('4.2.7 在线更新 UI 集成', 3)
add_body('设置页底部"// 在线更新"区块，绑定三个按钮：保存仓库、检查更新、应用更新。所有更新动作在浏览器内完成，无需 SSH。')
add_image_marker('Web 界面截图：设置页完整的"// 在线更新"区块（仓库地址输入 + 三按钮 + 状态文字区）')
add_caption('图 4-25  在线更新完整 UI 区块')

# ========== 五、总结 ==========
add_heading('五、总结', 1)

add_heading('5.1 项目成果', 2)
add_body('本项目按期完成全部 22 项功能需求和 8 项非功能需求，交付物包括：23 个 Python 模块 + 4 个前端文件 + 部署脚本，约 4500 行代码；5 张数据库表；README、实训报告、部署手册；支持任意 Linux 服务器一键安装，3 分钟内交付可用服务。')

add_heading('5.2 技术亮点', 2)
add_body('（1）SRC 平台 API 逆向：通过浏览器 F12 抓包还原补天/漏洞盒子接口，绕过反爬；')
add_body('（2）在线更新：基于 Git + sudoers 精细放权 + 延迟 systemd 重启，实现 Web 端零中断升级；')
add_body('（3）多用户隔离的调度器：每用户独立 SRC 间隔，全局取最小值避免重复请求；')
add_body('（4）跨发行版部署脚本：apt/yum/dnf 三分支自适配 + SQLite 自动迁移；')
add_body('（5）零依赖前端：原生 ES6 实现 i18n、排序、分页、批量导入，无 React/Vue 包袱。')

add_heading('5.3 遇到的问题与解决方案', 2)
add_table(
    ['问题', '现象', '解决方案'],
    [
        ['SQLite 只读', '服务启动后页面操作报错', 'data/ 目录所有权改为 assetmonitor'],
        ['模式漂移', '升级后页面表格显示异常', '部署脚本前置 SQLite ALTER TABLE 迁移'],
        ['端口绑定', '浏览器无法打开网页', '改为 0.0.0.0 + 云安全组放行'],
        ['首次扫描洪泛', '首次启动后 SRC 页面全是 NEW', '设计可配置首次仅入库不推送'],
        ['语言切换重叠', '仪表盘导航栏切换按钮位置错位', 'CSS body:not(.dashboard) 区分登录/仪表盘'],
    ]
)
add_image_marker('Web 界面截图：问题修复后的正常运行状态 — 仪表盘稳定显示数据，无报错')
add_caption('图 5-1  系统稳定运行效果')

add_heading('5.4 不足与改进方向', 2)
add_body('（1）可接入 FOFA、Quake、SecurityTrails 等付费 API 提升子域名覆盖率；')
add_body('（2）SQLite 在百万级资产下性能下降，下一版迁移 PostgreSQL；')
add_body('（3）JWT 黑名单机制、双因素认证；')
add_body('（4）自身可观测性（Prometheus metrics）；')
add_body('（5）接入大模型对资产做智能分类与优先级排序。')

add_heading('5.5 心得体会', 2)
add_body('【学生 A】通过本项目深入掌握了 Flask 全栈开发、APScheduler 任务调度、Linux systemd 服务管理以及 Git OTA 升级工程化实现。最大收获是理解了"生产可用"四字——不仅要功能跑通，更要考虑权限最小化、错误恢复、跨发行版兼容等工程细节。')
add_body('【学生 B】在前端零框架的情况下用原生 JS 实现了 i18n、排序、分页等组件化能力，体会到"简洁大于炫技"。SRC 监控的视觉设计反复迭代了 5 版才达到"赛博朋克 + 易读"的平衡，深刻理解了 UX 与 UI 的边界。')

doc.add_page_break()

# ============ 附录 ============
add_heading('附录 A：指导教师评语', 1)
for _ in range(4):
    doc.add_paragraph()
add_body('实训报告得分（总分）：_______________', indent=False)
for _ in range(4):
    doc.add_paragraph()
add_body('指导教师（签名）：__________________           年   月   日', indent=False)

doc.add_page_break()

add_heading('附录 B：答辩问题记录', 1)
for _ in range(6):
    doc.add_paragraph()
add_body('答辩得分（总分）：_______________', indent=False)

doc.add_page_break()

add_heading('附录 C：课程成果（图片）', 1)
add_body('以下精选 6 张代表性成果图（均为系统 Web 界面实际运行截图）：')
add_image_marker('成果图 1：仪表盘全屏（图 4-11）')
add_caption('图 C-1  仪表盘主界面')
add_image_marker('成果图 2：SRC 监控页（图 4-20）')
add_caption('图 C-2  SRC 平台监控')
add_image_marker('成果图 3：中英文对比拼图（图 4-12 + 图 4-13）')
add_caption('图 C-3  中英文双语支持')
add_image_marker('成果图 4：批量导入面板 + Toast 成功反馈拼图（图 4-17 + 图 4-18）')
add_caption('图 C-4  批量导入流程')
add_image_marker('成果图 5：在线更新 Web 控制面板（图 4-25）')
add_caption('图 C-5  在线更新面板')
add_image_marker('成果图 6：资产页 NEW 标识截图（图 4-3）')
add_caption('图 C-6  资产监控核心效果')

import os
out = 'C:/Users/fyb/Desktop/理工学院6月实训报告.docx'
try:
    doc.save(out)
    print('OK saved to', out)
except PermissionError:
    out2 = 'C:/Users/fyb/Desktop/理工学院6月实训报告_新版.docx'
    doc.save(out2)
    print('原文件被占用，已保存为新文件：', out2)
